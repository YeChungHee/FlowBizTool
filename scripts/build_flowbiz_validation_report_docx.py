from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


BASE_DIR = Path(__file__).resolve().parents[1]
OUT_PATH = BASE_DIR / "FlowBiz_ultra_고도화계획서_검증보고서_20260427.docx"
DETAIL_MD_PATH = BASE_DIR / "docs" / "flowbiz_ultra_pipeline_logic_validation_20260427.md"
REPORT_SERIAL = "FBU-VAL-0006"
REGISTRY_PATH = BASE_DIR / "docs" / "flowbiz_ultra_validation_report_registry.md"

FONT_NAME = "AppleGothic"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_RED = "FCE4D6"
LIGHT_YELLOW = "FFF2CC"
GRAY = "F2F2F2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    cell.width = Twips(width_twips)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for col_idx, width in enumerate(widths):
        for cell in table.columns[col_idx].cells:
            set_cell_width(cell, width)


def set_table_cell_margins(table, top=100, start=120, bottom=100, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10)

    for style_name, size, color in [
        ("Title", 22, BLUE),
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12, BLUE),
        ("Heading 3", 10.5, BLUE),
    ]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_paragraph(doc: Document, text: str = "", *, style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_font(r)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], header_fill: str = LIGHT_BLUE) -> None:
    # Artifact-tool currently renders some python-docx tables with overly narrow
    # columns. Keep the validation report robust by representing tabular facts as
    # structured paragraph blocks instead of OOXML tables.
    for row_values in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Pt(8)
        first = p.add_run(str(row_values[0]))
        set_font(first, size=9.4, bold=True, color=BLUE)
        for header, value in zip(headers[1:], row_values[1:]):
            mid = p.add_run(f"  |  {header}: ")
            set_font(mid, size=9.2, bold=True)
            val = p.add_run(str(value))
            set_font(val, size=9.2)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_YELLOW) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p_pr = p._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    r = p.add_run(title)
    set_font(r, size=10, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.left_indent = Pt(8)
    r2 = p2.add_run(body)
    set_font(r2, size=9.2)
    doc.add_paragraph()


def add_markdown_detail_appendix(doc: Document, md_path: Path) -> None:
    if not md_path.is_file():
        return

    doc.add_page_break()
    doc.add_heading("9. 상세 검증 보고서 전문", level=1)
    add_paragraph(
        doc,
        "아래 내용은 데이터 입력 파이프라인, 평가출력 파이프라인, Review Findings 5개, 개선 계획을 포함한 상세 검증 보고서 전문이다.",
    )

    in_code = False
    for raw_line in md_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if not stripped:
            continue

        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(line)
            set_font(r, size=8.5, color="444444")
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue

        if stripped.startswith("|"):
            normalized = stripped.strip("|").replace("|", "  |  ")
            if set(normalized.replace(" ", "").replace("-", "").replace(":", "")) == set():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(normalized)
            set_font(r, size=8.8)
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(stripped[2:].strip())
            set_font(r, size=9.2)
            continue

        number_match = stripped.split(".", 1)
        if len(number_match) == 2 and number_match[0].isdigit() and number_match[1].startswith(" "):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(number_match[1].strip())
            set_font(r, size=9.2)
            continue

        add_paragraph(doc, stripped)


def build_doc() -> Document:
    doc = Document()
    style_doc(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("FlowBiz_ultra\n고도화 계획서 검증 보고서")
    set_font(r, size=23, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"일련번호: {REPORT_SERIAL}\n"
        "검증 기준일: 2026년 4월 27일\n"
        "대상 문서: FlowBiz_ultra_고도화계획서_20260427.docx\n"
        "작성 위치: /Users/appler/Documents/COTEX/FlowBiz_ultra\n"
        "누적관리 대장: docs/flowbiz_ultra_validation_report_registry.md"
    )
    set_font(r, size=10)

    add_callout(
        doc,
        "종합 판정",
        "계획서의 기술 방향은 유효하다. 단, 실행은 SourceQuality -> registry 재정규화 -> active framework/promote -> 회귀검증 순서로 진행해야 하며, 원본 DOCX는 스타일 중복과 표 레이아웃 문제를 수정한 뒤 최종 배포해야 한다.",
        fill=LIGHT_GREEN,
    )
    doc.add_page_break()

    doc.add_heading("1. 검증 개요", level=1)
    add_table(
        doc,
        ["항목", "검증 결과", "판정"],
        [
            ["일련번호", f"{REPORT_SERIAL} / 다음 번호 FBU-VAL-0007", "누적관리 적용"],
            ["계획 내용", "P0/P1/P2 리스크와 4단계 고도화 로드맵은 현재 코드 진단과 일치한다.", "승인 가능"],
            ["실행 순서", "SourceQuality와 학습 적격 게이트가 선행되어야 한다.", "조건부 승인"],
            ["현재 구현 상태", "SourceQuality, active framework, promote, pytest 체계는 아직 미구현이다.", "추가 구현 필요"],
            ["DOCX 구조", "Heading1/2/3 styleId 중복으로 artifact-tool 원본 렌더링이 실패했다.", "수정 필요"],
            ["문서 가독성", "임시 복구본 렌더링 기준 일부 표가 좁게 배치되어 세로 글자처럼 보인다.", "수정 필요"],
        ],
        [2100, 5600, 1700],
    )

    doc.add_heading("2. 코드 검증 결과", level=1)
    add_table(
        doc,
        ["검증 항목", "결과", "근거"],
        [
            ["컴파일", "통과", "app.py, engine.py, report_extractors.py, bizaipro_learning.py 등 py_compile 성공"],
            ["학습 상태", "후보 7건 / 업데이트 적격 5건", "누적 update weight 4.30, 업데이트 조건 미충족"],
            ["업데이트 실행", "미생성", "적격 10건과 가중치 7.5 조건 미달로 update_generated=false"],
            ["Notion 상담보고서", "일부 성공", "씨랩, 쿳션 상담보고서 official API 본문 추출 성공"],
            ["Notion 심사보고서", "일부 실패", "씨랩, 쿳션 심사보고서는 integration 공유 권한 문제로 본문 0자"],
            ["테스트 체계", "미구성", "tests 디렉터리, pytest.ini, requirements.txt, pyproject.toml 없음"],
        ],
        [2300, 2300, 4800],
    )

    doc.add_heading("3. 핵심 리스크", level=1)
    add_table(
        doc,
        ["우선순위", "이슈", "영향", "권고"],
        [
            ["P0", "업데이트 산출물이 실제 평가엔진에 자동 적용되지 않는다.", "버전명만 바뀌고 웹 평가 결과가 변하지 않을 수 있다.", "framework snapshot, active version registry, promote 명령 추가"],
            ["P0", "자료 품질이 아니라 파일명/URL 존재 여부로 학습 적격을 계산한다.", "본문 추출 실패 자료도 update weight를 받을 수 있다.", "SourceQuality 기반 usable_for_update 적용"],
            ["P1", "CLI 학습 루프와 웹 학습 루프의 적격 기준이 다르다.", "혼재된 품질 기준의 학습 케이스가 registry에 쌓일 수 있다.", "공통 learning status 모듈로 통합"],
            ["P1", "문서 추출 결과가 세부 평가항목으로 충분히 매핑되지 않는다.", "상담/심사 근거가 structure bonus 중심으로만 반영된다.", "subfactor mapping과 evidence 저장"],
            ["P2", "결과 검증 자동화가 부족하다.", "오류 있는 업데이트가 promote될 수 있다.", "pytest와 shadow evaluation 게이트 도입"],
        ],
        [1200, 3300, 2500, 2400],
        header_fill=LIGHT_RED,
    )

    doc.add_heading("4. 계획서 실행 가능성", level=1)
    add_table(
        doc,
        ["단계", "계획 항목", "실행 판정", "선행 조건"],
        [
            ["1", "학습 적격 게이트 정비", "즉시 착수 가능", "SourceQuality helper, parser quality 반환, registry 백업"],
            ["2", "평가 입력 변환 고도화", "1단계 후 착수", "품질 기준, evidence/confidence 저장 구조"],
            ["3", "업데이트 적용 경로 구축", "registry 재정규화 후 착수", "정제된 학습 케이스와 shadow evaluation 기준"],
            ["4", "결과 검증 자동화", "스캐폴드 병행 가능", "pytest 설정, fixture, promote 차단 규칙"],
        ],
        [900, 2900, 2200, 3400],
    )

    doc.add_heading("5. 보완 권고", level=1)
    add_bullets(
        doc,
        [
            "SourceQuality는 usable_for_evaluation과 usable_for_update를 분리한다.",
            "기존 registry 7건은 재정규화 전후 비교표를 별도 산출물로 남긴다.",
            "Notion 권한 오류는 UI에서 자료 있음이 아니라 본문 미확인으로 표시한다.",
            "promote 기준에는 decision 변화율, 한도 변화율, P0 오류 0건을 포함한다.",
            "각 단계 완료 기준에 실행 명령과 산출물 경로를 명시한다.",
        ],
    )

    doc.add_heading("6. 즉시 실행 순서", level=1)
    add_table(
        doc,
        ["순서", "작업", "대상"],
        [
            ["1", "SourceQuality 데이터 구조와 품질 판정 helper 추가", "app.py 또는 공통 모듈"],
            ["2", "FlowScore, Notion, supporting/additional parser에 quality 반환 추가", "report_extractors.py, app.py"],
            ["3", "learning_material_components를 품질 기반으로 전환", "app.py, bizaipro_learning.py"],
            ["4", "기존 registry 7건 백업 및 재정규화", "data/bizaipro_learning_registry.json"],
            ["5", "재정규화 전후 비교 리포트 생성", "docs 또는 outputs"],
            ["6", "active framework/promote 구현", "bizaipro_learning.py, engine.py, app.py"],
            ["7", "pytest fixture와 회귀검증 명령 추가", "tests/, pytest.ini"],
        ],
        [900, 5600, 2900],
        header_fill=LIGHT_GREEN,
    )

    doc.add_heading("7. Go / No-Go 판정", level=1)
    add_table(
        doc,
        ["항목", "판정", "비고"],
        [
            ["계획서 내용 기준 고도화 착수", "Go", "기술 방향은 현재 코드 진단과 일치"],
            ["1단계 SourceQuality 구현", "Go", "가장 먼저 해결해야 할 P0 리스크"],
            ["promote 구현 선착수", "No-Go", "품질 기준 정비 전에는 후보 데이터 신뢰도가 낮음"],
            ["원본 DOCX 최종 배포", "No-Go", "중복 스타일과 표 레이아웃 수정 필요"],
            ["새 검증보고서 배포", "Go", "본 문서는 새 스타일로 생성되어 렌더 검증 예정"],
        ],
        [3200, 1600, 4600],
    )

    doc.add_heading("8. 검증 명령", level=1)
    add_paragraph(doc, "이번 검증에서 사용한 핵심 명령은 다음과 같다.")
    add_table(
        doc,
        ["목적", "명령"],
        [
            ["컴파일 확인", "python3 -m py_compile app.py engine.py report_extractors.py bizaipro_learning.py external_apis.py proposal_generator.py"],
            ["학습 상태", "python3 bizaipro_learning.py status"],
            ["업데이트 시도", "python3 bizaipro_learning.py update"],
            ["DOCX 렌더", "render_docx.py --renderer artifact-tool"],
            ["DOCX 접근성", "a11y_audit.py FlowBiz_ultra_고도화계획서_20260427.docx"],
        ],
        [2200, 7200],
    )

    doc.add_heading("9. 누적관리 정보", level=1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["현재 보고서 일련번호", REPORT_SERIAL],
            ["누적관리 대장", str(REGISTRY_PATH.relative_to(BASE_DIR))],
            ["다음 신규 검증번호", "FBU-VAL-0007"],
            ["수정본 번호 규칙", f"{REPORT_SERIAL}-R1, {REPORT_SERIAL}-R2"],
        ],
        [2600, 6800],
        header_fill=LIGHT_GREEN,
    )

    add_markdown_detail_appendix(doc, DETAIL_MD_PATH)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("COTEX / FlowBiz_ultra validation report / 2026-04-27")
    set_font(run, size=8, color="666666")
    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
